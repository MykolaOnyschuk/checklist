#!/usr/bin/env python3
"""
Maintenance Trail Report Generator
Proactive Safety-First Home Maintenance — Utah Market
"""

import datetime
import os
import sys

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ROW_HEIGHT_RULE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Error: python-docx not installed. Run: pip3 install python-docx")
    sys.exit(1)


# ─── COLOR CONSTANTS (OOXML hex fill strings) ────────────────────────────────
HEADER_HEX   = "1A3A5C"   # dark navy  — main section headers
SUBHDR_HEX   = "E3F2FD"   # pale blue  — table column labels
PASS_HEX     = "C8E6C9"   # pale green
FAIL_HEX     = "FFCDD2"   # pale red
ATTN_HEX     = "FFF9C4"   # pale yellow
SKIP_HEX     = "F5F5F5"   # light gray — N/A items
PHOTO_HEX    = "E0E0E0"   # mid gray   — photo placeholder boxes
PROP_HEX     = "FFF3E0"   # pale orange — proposals row

DARK_BLUE  = RGBColor(0x1A, 0x3A, 0x5C)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
MID_GRAY   = RGBColor(0x61, 0x61, 0x61)
LIGHT_GRAY = RGBColor(0x75, 0x75, 0x75)


# ─── CHECKLISTS ───────────────────────────────────────────────────────────────

STR_SENTRY = [
    {
        "section": "Life Safety",
        "items": [
            "Smoke detector tested — all units beep on test",
            "CO detector tested — all units respond correctly",
            "Fire extinguisher — present, charged, pin intact",
            "Emergency egress path — clear and unobstructed",
        ],
    },
    {
        "section": "Water & Leak Check",
        "items": [
            "Under kitchen sink — no leaks, moisture, or mold",
            "Under bathroom sink(s) — dry, no corrosion on supply lines",
            "Around toilet base(s) — no seepage or staining",
            "Water heater area — no rust staining, no pooling",
            "Washer supply hoses (if applicable) — secure and dry",
        ],
    },
    {
        "section": "Smart Access & Security",
        "items": [
            "Smart lock battery level — adequate (>20%)",
            "Smart lock function test — locks and unlocks correctly",
            "Guest access code verified and operational",
            "Backup key / lockbox confirmed accessible (if applicable)",
        ],
    },
    {
        "section": "HVAC & Climate",
        "items": [
            "HVAC filter — visual condition acceptable",
            "Thermostat operation — both heating and cooling mode",
        ],
    },
    {
        "section": "General Property Safety",
        "items": [
            "Interior lighting — all fixtures operational",
            "Exterior entry lighting — functional",
            "Trip hazards — none identified in main living areas",
            "Stair handrail — secure and stable",
        ],
    },
]

STR_GUARDIAN_EXTRA = [
    {
        "section": "Appliance Check",
        "items": [
            "Dishwasher — completes cycle, no leaks at door seal",
            "Refrigerator — temperature adequate (35–38 °F)",
            "Microwave — operational",
            "Range / oven — all burners and oven element functional",
            "Washer — runs cycle, drains correctly",
            "Dryer — operates, lint trap clean, vent not restricted",
        ],
    },
    {
        "section": "HVAC Deep Check",
        "items": [
            "HVAC filter — replaced or scheduled for replacement",
            "Airflow adequate at all supply vents",
            "Condensate drain line — clear, no overflow sign",
            "Outdoor unit visual — no damage, debris cleared 12\" around unit",
        ],
    },
    {
        "section": "Exterior Perimeter",
        "items": [
            "All exterior lights — operational",
            "Entry door(s) — hardware secure, weatherstrip intact",
            "Window locks — all operable and engaging correctly",
            "Gutters / downspouts — no visible blockage or damage",
            "Driveway / walkway — no significant cracks or trip hazards",
        ],
    },
]

EDUCATIONAL = [
    {
        "section": "Life Safety Compliance  [CRITICAL — 1-hr response if fail]",
        "items": [
            "Fire alarm panel — normal status, no trouble or alarm lights",
            "Manual pull stations — accessible, no physical damage",
            "Emergency lighting — 30-second test passed, all units illuminate",
            "Exit signs — all illuminated (primary power + battery backup)",
            "Fire extinguishers — inspection tags current, gauges in green zone",
            "Sprinkler heads — visible, no obstructions within 18\"",
        ],
    },
    {
        "section": "ADA Accessibility & Entry  [CRITICAL — 1-hr response if fail]",
        "items": [
            "Main entrance door closer — smooth operation, latches fully",
            "Door threshold — no raised edges, ADA compliant",
            "Accessible parking signage — visible and compliant",
            "Accessible path from parking to entrance — clear of obstructions",
            "ADA restroom grab bars — secure, clearances maintained",
        ],
    },
    {
        "section": "HVAC — Predictive Maintenance  [URGENT — 4-hr response if fail]",
        "items": [
            "Air filters — condition checked, replacement date recorded",
            "Thermostat — calibration verified, programming correct",
            "Airflow at classroom supply vents — adequate and consistent",
            "Condensate pan — dry, drain line clear and flowing",
            "Belt / motor (if applicable) — no unusual noise, no visible wear",
        ],
    },
    {
        "section": "Plumbing  [URGENT — 4-hr response if fail]",
        "items": [
            "Restroom faucets — no drips, adequate water pressure",
            "Toilets — flush and fill cycle operating correctly",
            "Under-sink visual — no leaks or moisture in all restrooms",
            "Water fountains — operational, adequate flow and drain",
            "No visible active plumbing leaks anywhere in facility",
        ],
    },
    {
        "section": "Electrical  [URGENT — 4-hr response if fail]",
        "items": [
            "GFCI outlets in wet areas — tested and reset correctly",
            "Electrical panel visual — no tripped breakers, no unusual odor",
            "Exterior lighting — all fixtures operational",
            "Classroom lighting — all fixtures functional, no burnt bulbs",
        ],
    },
    {
        "section": "Security & Access Control  [CRITICAL — 1-hr response if fail]",
        "items": [
            "Exterior door locks / deadbolts — all operational",
            "Classroom door locks — engage and release correctly",
            "Security camera system — all cameras showing live feed",
            "Intercom / buzzer entry system — operational",
        ],
    },
    {
        "section": "General Safety & Facility  [ROUTINE — 48-hr response]",
        "items": [
            "Slip / trip hazards — none identified in hallways or classrooms",
            "Flooring condition — no damaged, lifted, or missing sections",
            "Ceiling tiles — no damaged, missing, or water-stained tiles",
            "Stair handrail integrity — secure, continuous full length",
            "Playground equipment (if applicable) — no sharp edges or loose hardware",
        ],
    },
]

MOVE_OUT = STR_SENTRY + STR_GUARDIAN_EXTRA


# ─── XML / FORMATTING HELPERS ─────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_table_borders(table, color="BDBDBD", size="4"):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                el = OxmlElement(f"w:{edge}")
                el.set(qn("w:val"), "single")
                el.set(qn("w:sz"), size)
                el.set(qn("w:space"), "0")
                el.set(qn("w:color"), color)
                tcBorders.append(el)
            tcPr.append(tcBorders)


def _row_height(row, height_twips):
    trPr = row._tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(height_twips))
    trHeight.set(qn("w:hRule"), "atLeast")
    trPr.append(trHeight)


def _cell_para(cell, text, size=9, bold=False, italic=False,
               align=WD_ALIGN_PARAGRAPH.LEFT, color=None,
               space_before=2, space_after=2):
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return run


def _status_hex(status):
    return {"PASS": PASS_HEX, "FAIL": FAIL_HEX, "ATTENTION": ATTN_HEX, "N/A": SKIP_HEX}.get(status, "FFFFFF")


# ─── DOCUMENT BUILDER ─────────────────────────────────────────────────────────

def build_report(plan_label, info, tech_name, checklist_results, work_done, time_used, proposals):
    doc = Document()

    # Page margins
    sec = doc.sections[0]
    sec.top_margin    = Cm(1.8)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin   = Cm(2.2)
    sec.right_margin  = Cm(2.2)

    # Default paragraph spacing
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9)

    # ── HEADER BANNER ─────────────────────────────────────────────────────────
    banner = doc.add_table(rows=1, cols=2)
    banner.autofit = False
    banner.columns[0].width = Inches(4.8)
    banner.columns[1].width = Inches(2.4)

    lc = banner.cell(0, 0)
    _set_cell_bg(lc, HEADER_HEX)
    lc.paragraphs[0].paragraph_format.space_before = Pt(8)
    r = lc.paragraphs[0].add_run("MAINTENANCE TRAIL REPORT")
    r.bold = True; r.font.size = Pt(17); r.font.color.rgb = WHITE
    p2 = lc.add_paragraph()
    p2.paragraph_format.space_after = Pt(8)
    r2 = p2.add_run("Proactive Safety-First Home Maintenance  ·  Utah")
    r2.font.size = Pt(8.5); r2.font.color.rgb = RGBColor(0xBB, 0xDE, 0xFB)

    rc = banner.cell(0, 1)
    _set_cell_bg(rc, HEADER_HEX)
    rp = rc.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_before = Pt(8)
    rr = rp.add_run("[  YOUR LOGO  ]")
    rr.italic = True; rr.font.size = Pt(9); rr.font.color.rgb = RGBColor(0x90, 0xCA, 0xF9)
    rp2 = rc.add_paragraph()
    rp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp2.paragraph_format.space_after = Pt(8)
    rr2 = rp2.add_run(f"Date: {info['date']}")
    rr2.font.size = Pt(8); rr2.font.color.rgb = RGBColor(0xBB, 0xDE, 0xFB)

    doc.add_paragraph()

    # ── CLIENT INFO TABLE ─────────────────────────────────────────────────────
    ci = doc.add_table(rows=4, cols=4)
    ci.autofit = False
    for col, w in zip(ci.columns, [1.2, 3.4, 1.4, 3.0]):
        col.width = Inches(w)
    _set_table_borders(ci, "C5CAE9", "4")

    rows_data = [
        ("Client / Owner",    info["client"],   "Property Address", info["address"]),
        ("Service Plan",      plan_label,        "Visit Type",       info["visit_type"]),
        ("Technician",        tech_name,         "Visit Date",       info["date"]),
        ("Report #",          info["report_num"],"Next Scheduled",   info["next_visit"]),
    ]
    for row_idx, (l1, v1, l2, v2) in enumerate(rows_data):
        row = ci.rows[row_idx]
        for col_idx, (txt, is_label) in enumerate([(l1,True),(v1,False),(l2,True),(v2,False)]):
            cell = row.cells[col_idx]
            _set_cell_bg(cell, SUBHDR_HEX if is_label else "FFFFFF")
            _cell_para(cell, txt, size=9, bold=is_label,
                       color=DARK_BLUE if is_label else RGBColor(0x21, 0x21, 0x21),
                       space_before=3, space_after=3)

    doc.add_paragraph()

    # ── CHECKLIST SECTIONS ────────────────────────────────────────────────────
    failed_items    = []
    attention_items = []

    for section_name, items_results in checklist_results:
        # Section header bar
        sh = doc.add_table(rows=1, cols=1)
        sh.autofit = True
        sc = sh.cell(0, 0)
        _set_cell_bg(sc, HEADER_HEX)
        sc.paragraphs[0].paragraph_format.space_before = Pt(4)
        sc.paragraphs[0].paragraph_format.space_after  = Pt(4)
        sr = sc.paragraphs[0].add_run(f"  {section_name.upper()}")
        sr.bold = True; sr.font.size = Pt(9.5); sr.font.color.rgb = WHITE

        # Checklist table: STATUS | ITEM | NOTES
        ct = doc.add_table(rows=len(items_results) + 1, cols=3)
        ct.autofit = False
        ct.columns[0].width = Inches(0.75)
        ct.columns[1].width = Inches(4.55)
        ct.columns[2].width = Inches(1.90)
        _set_table_borders(ct, "CFD8DC", "4")

        for hc, ht in zip(ct.rows[0].cells, ["STATUS", "INSPECTION ITEM", "NOTES"]):
            _set_cell_bg(hc, SUBHDR_HEX)
            _cell_para(hc, ht, size=8, bold=True, color=DARK_BLUE)

        for r_idx, (item, status, note) in enumerate(items_results):
            row = ct.rows[r_idx + 1]
            sc2, ic, nc = row.cells

            _set_cell_bg(sc2, _status_hex(status))
            sp = sc2.paragraphs[0]
            sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sp.paragraph_format.space_before = Pt(3)
            sp.paragraph_format.space_after  = Pt(3)
            sr2 = sp.add_run(status)
            sr2.bold = True; sr2.font.size = Pt(8)

            _cell_para(ic, item, size=9)
            _cell_para(nc, note, size=8, color=MID_GRAY)

            if status == "FAIL":
                failed_items.append((item, note))
            elif status == "ATTENTION":
                attention_items.append((item, note))

        doc.add_paragraph()

    # ── ISSUES & PHOTO DOCUMENTATION ─────────────────────────────────────────
    all_issues = [("FAIL", i, n) for i, n in failed_items] + \
                 [("ATTENTION", i, n) for i, n in attention_items]

    if all_issues:
        doc.add_page_break()

        ih = doc.add_table(rows=1, cols=1)
        ihc = ih.cell(0, 0)
        _set_cell_bg(ihc, "B71C1C")
        ihc.paragraphs[0].paragraph_format.space_before = Pt(5)
        ihc.paragraphs[0].paragraph_format.space_after  = Pt(5)
        ihr = ihc.paragraphs[0].add_run("  ISSUES REQUIRING ACTION — PHOTO DOCUMENTATION")
        ihr.bold = True; ihr.font.size = Pt(11); ihr.font.color.rgb = WHITE

        doc.add_paragraph()

        for idx, (severity, item, note) in enumerate(all_issues, 1):
            # Issue label
            it = doc.add_table(rows=1, cols=1)
            itc = it.cell(0, 0)
            _set_cell_bg(itc, FAIL_HEX if severity == "FAIL" else ATTN_HEX)
            itc.paragraphs[0].paragraph_format.space_before = Pt(4)
            itr = itc.paragraphs[0].add_run(f"  Issue #{idx}  [{severity}]   {item}")
            itr.bold = True; itr.font.size = Pt(10)

            if note:
                np = itc.add_paragraph()
                np.paragraph_format.space_after = Pt(5)
                nr = np.add_run(f"  Technician Note: {note}")
                nr.font.size = Pt(9); nr.italic = True

            doc.add_paragraph()

            # Two-column photo placeholder table
            pt = doc.add_table(rows=1, cols=2)
            pt.autofit = False
            pt.columns[0].width = Inches(3.55)
            pt.columns[1].width = Inches(3.55)
            _set_table_borders(pt, "9E9E9E", "6")

            for ph_idx, ph_label in enumerate(
                ["BEFORE / Issue Photo\n(click to insert image)",
                 "AFTER / Resolution Photo\n(click to insert image)"]
            ):
                phc = pt.cell(0, ph_idx)
                _set_cell_bg(phc, PHOTO_HEX)
                _row_height(pt.rows[0], 1800)   # ~1.25 inches minimum
                p = phc.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(24)
                p.paragraph_format.space_after  = Pt(24)
                pr = p.add_run(f"[ {ph_label} ]")
                pr.font.size = Pt(9.5)
                pr.bold = True
                pr.font.color.rgb = MID_GRAY

            doc.add_paragraph()

    # ── WORK COMPLETED ────────────────────────────────────────────────────────
    doc.add_page_break()

    wh = doc.add_table(rows=1, cols=1)
    whc = wh.cell(0, 0)
    _set_cell_bg(whc, HEADER_HEX)
    whc.paragraphs[0].paragraph_format.space_before = Pt(5)
    whc.paragraphs[0].paragraph_format.space_after  = Pt(5)
    whr = whc.paragraphs[0].add_run("  WORK COMPLETED THIS VISIT")
    whr.bold = True; whr.font.size = Pt(11); whr.font.color.rgb = WHITE

    doc.add_paragraph()

    wt = doc.add_table(rows=2, cols=2)
    wt.autofit = False
    wt.columns[0].width = Inches(5.1)
    wt.columns[1].width = Inches(2.0)
    _set_table_borders(wt, "CFD8DC", "4")

    for hc, ht in zip(wt.rows[0].cells, ["TASKS PERFORMED", "LABOR USED"]):
        _set_cell_bg(hc, SUBHDR_HEX)
        _cell_para(hc, ht, size=9, bold=True, color=DARK_BLUE)

    _cell_para(wt.cell(1, 0), work_done or "No repair work required this visit.", size=9,
               space_before=5, space_after=5)
    _cell_para(wt.cell(1, 1), time_used or "—", size=9, space_before=5, space_after=5)

    doc.add_paragraph()

    # ── PROPOSALS TABLE ───────────────────────────────────────────────────────
    if proposals:
        proph = doc.add_table(rows=1, cols=1)
        prophc = proph.cell(0, 0)
        _set_cell_bg(prophc, "E65100")
        prophc.paragraphs[0].paragraph_format.space_before = Pt(5)
        prophc.paragraphs[0].paragraph_format.space_after  = Pt(5)
        prophr = prophc.paragraphs[0].add_run("  FIXED-PRICE PROPOSALS FOR CLIENT APPROVAL")
        prophr.bold = True; prophr.font.size = Pt(11); prophr.font.color.rgb = WHITE

        doc.add_paragraph()

        propt = doc.add_table(rows=len(proposals) + 1, cols=5)
        propt.autofit = False
        for col, w in zip(propt.columns, [3.2, 0.9, 0.9, 0.9, 1.2]):
            col.width = Inches(w)
        _set_table_borders(propt, "CFD8DC", "4")

        for hc, ht in zip(propt.rows[0].cells,
                          ["REPAIR DESCRIPTION", "LABOR ($)", "PARTS ($)", "TRIP FEE ($)", "TOTAL ($)"]):
            _set_cell_bg(hc, SUBHDR_HEX)
            _cell_para(hc, ht, size=8.5, bold=True, color=DARK_BLUE)

        for row_idx, (desc, labor, parts, trip, total) in enumerate(proposals, 1):
            row = propt.rows[row_idx]
            for cell, val in zip(row.cells, [desc, labor, parts, trip, total]):
                _set_cell_bg(cell, PROP_HEX)
                _cell_para(cell, str(val), size=9)

        doc.add_paragraph()
        np = doc.add_paragraph()
        nr = np.add_run(
            "CLIENT APPROVAL REQUIRED before scheduling any proposed repair. "
            "Reply to this report or contact your technician to authorize work."
        )
        nr.font.size = Pt(8); nr.italic = True; nr.font.color.rgb = LIGHT_GRAY

        doc.add_paragraph()

    # ── SIGNATURE BLOCK ───────────────────────────────────────────────────────
    sigt = doc.add_table(rows=2, cols=3)
    sigt.autofit = False
    for col, w in zip(sigt.columns, [2.4, 2.4, 2.4]):
        col.width = Inches(w)
    _set_table_borders(sigt, "90A4AE", "4")

    for hc, ht in zip(sigt.rows[0].cells,
                      ["Technician Signature", "Client Signature / Approval", "Date Signed"]):
        _set_cell_bg(hc, SUBHDR_HEX)
        _cell_para(hc, ht, size=8.5, bold=True, color=DARK_BLUE)

    for bc, val in zip(sigt.rows[1].cells, [tech_name, "", info["date"]]):
        bc.paragraphs[0].paragraph_format.space_before = Pt(22)
        bc.paragraphs[0].paragraph_format.space_after  = Pt(6)
        r = bc.paragraphs[0].add_run(val)
        r.font.size = Pt(9)

    # ── FOOTER ────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(
        "This Maintenance Trail report is part of your permanent property record. "
        "Retain for warranty, insurance, and compliance purposes.  ·  Utah Licensed & Insured."
    )
    fr.font.size = Pt(7); fr.italic = True; fr.font.color.rgb = LIGHT_GRAY

    return doc


# ─── CLI ──────────────────────────────────────────────────────────────────────

PLAN_MAP = {
    "1": ("STR — Sentry Plan  ($200/mo · Safety Sweep)", STR_SENTRY, "Scheduled"),
    "2": ("STR — Guardian Plan  ($300/mo · Deep Dive)",  STR_SENTRY + STR_GUARDIAN_EXTRA, "Scheduled"),
    "3": ("Educational Facility Inspection",              EDUCATIONAL, "Scheduled"),
    "4": ("Move-Out / Vacant Property Inspection",        MOVE_OUT, "Move-Out"),
}


def prompt_checklist(sections):
    all_results = []
    print("\n  For each item enter:  P = Pass   F = Fail   A = Attention   S = Skip/N/A")
    for sec in sections:
        print(f"\n  ▶ {sec['section'].upper()}")
        sec_results = []
        for item in sec["items"]:
            while True:
                choice = input(f"\n    {item[:70]}\n    > ").strip().upper()
                if choice in ("P", "F", "A", "S", ""):
                    break
                print("    Enter P, F, A, or S")
            status = {"P": "PASS", "F": "FAIL", "A": "ATTENTION", "S": "N/A", "": "N/A"}[choice]
            note = ""
            if status in ("FAIL", "ATTENTION"):
                note = input("    Note (describe issue): ").strip()
            sec_results.append((item, status, note))
        all_results.append((sec["section"], sec_results))
    return all_results


def prompt_proposals():
    proposals = []
    print("\n── FIXED-PRICE PROPOSALS ──────────────────────────────────────────")
    print("  Add a proposal for any repair that exceeds the monthly labor allowance.")
    while True:
        add = input("\n  Add a proposal? (y/n): ").strip().lower()
        if add != "y":
            break
        desc  = input("  Repair description: ").strip()
        labor = input("  Labor cost ($): ").strip()
        parts = input("  Parts / materials cost ($): ").strip()
        trip  = input("  Trip / premium fee ($, press Enter for 0): ").strip() or "0"
        try:
            total = float(labor or 0) + float(parts or 0) + float(trip or 0)
        except ValueError:
            total = 0.0
        proposals.append((desc, f"${labor}", f"${parts}", f"${trip}", f"${total:.2f}"))
    return proposals


def main():
    print("\n" + "=" * 62)
    print("  MAINTENANCE TRAIL REPORT GENERATOR")
    print("  Proactive Safety-First Home Maintenance — Utah Market")
    print("=" * 62)

    print("\nSelect Report Type:")
    for key, (label, _, _) in PLAN_MAP.items():
        print(f"  {key}. {label}")

    while True:
        choice = input("\nEnter choice (1–4): ").strip()
        if choice in PLAN_MAP:
            break

    plan_label, checklist, visit_type = PLAN_MAP[choice]
    today = str(datetime.date.today())

    print("\n── CLIENT & PROPERTY INFORMATION ──────────────────────────────────")
    info = {
        "client":     input("  Client / Owner Name: ").strip(),
        "address":    input("  Property Address: ").strip(),
        "date":       input(f"  Visit Date (Enter for today {today}): ").strip() or today,
        "next_visit": input("  Next Scheduled Visit Date: ").strip() or "TBD",
        "report_num": input("  Report # (e.g. 2025-001): ").strip() or "—",
        "visit_type": visit_type,
    }
    tech_name  = input("  Technician Name: ").strip()

    print("\n── INSPECTION CHECKLIST ────────────────────────────────────────────")
    checklist_results = prompt_checklist(checklist)

    print("\n── WORK COMPLETED ──────────────────────────────────────────────────")
    print("  Describe all repair work performed. (Blank line to finish.)")
    lines = []
    while True:
        line = input("  > ")
        if not line and lines:
            break
        lines.append(line)
    work_done  = "\n".join(lines).strip()
    time_used  = input("  Labor time used (e.g. '1.5 hrs of 2-hr allowance'): ").strip()

    proposals = prompt_proposals()

    print("\n  Generating report...")
    doc = build_report(plan_label, info, tech_name, checklist_results, work_done, time_used, proposals)

    out_dir = os.path.expanduser("~/Reports")
    os.makedirs(out_dir, exist_ok=True)

    safe_addr = "".join(c for c in info["address"] if c.isalnum() or c in " -_")[:35].strip().replace(" ", "_")
    safe_date = info["date"].replace("/", "-")
    filename  = f"MaintenanceTrail_{safe_addr}_{safe_date}.docx"
    filepath  = os.path.join(out_dir, filename)
    doc.save(filepath)

    print(f"\n  Report saved:")
    print(f"  {filepath}")
    print("\n  To convert to PDF:")
    print("  • Mac: open the file in Word → File → Print → Save as PDF")
    print("  • Or: open in Word → File → Export → PDF")
    print()


if __name__ == "__main__":
    main()
